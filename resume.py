import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Load JSON ──────────────────────────────────────────────
with open(sys.argv[1], "r", encoding="utf-8") as f:
    data = json.load(f)

TEMPLATE_PATH = "template.docx"
OUTPUT_FILENAME = sys.argv[2] if len(sys.argv) > 2 else "output_resume.docx"
OUTPUT_DIR = "output"

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Final path
OUTPUT_PATH = os.path.join(OUTPUT_DIR, OUTPUT_FILENAME)

# ── Colors ─────────────────────────────────────────────────
ACCENT = RGBColor(0x1A, 0x3C, 0x5E)
GRAY   = RGBColor(0x55, 0x55, 0x55)
LGRAY  = RGBColor(0x88, 0x88, 0x88)

# ── Helpers ────────────────────────────────────────────────
def set_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)

def set_indent(para, left=360, hanging=260):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),    str(left))
    ind.set(qn("w:hanging"), str(hanging))
    pPr.append(ind)

def bullet_para(doc, text):
    para = doc.add_paragraph()
    set_spacing(para, before=20, after=20)
    set_indent(para, left=360, hanging=260)
    rb = para.add_run("\u2022  ")
    rb.font.size      = Pt(9)
    rb.font.color.rgb = GRAY
    rb.font.name      = "Arial"
    rt = para.add_run(text)
    rt.font.size      = Pt(9)
    rt.font.color.rgb = GRAY
    rt.font.name      = "Arial"
    return para

def job_header_para(doc, title, company, city, dates):
    """
    Line 1: Job Title  —  Company, City          (bold navy)
    Line 2: dates                                 (italic gray, indented)
    Two separate paragraphs — avoids tab overflow on long company names.
    """
    # Line 1 — title + company
    p1 = doc.add_paragraph()
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = ACCENT
    r1.font.name = "Arial"
    r2 = p1.add_run(f"  \u2014  {company}, {city}")
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    r2.font.name = "Arial"
    set_spacing(p1, before=80, after=0)

    # Line 2 — dates only, indented to sit under company name
    p2 = doc.add_paragraph()
    rd = p2.add_run(dates)
    rd.italic = True
    rd.font.size = Pt(9)
    rd.font.color.rgb = LGRAY
    rd.font.name = "Arial"
    set_spacing(p2, before=0, after=10)

    return p1, p2

def project_header_para(doc, name, tech_stack, date):
    """
    Line 1: Project Name  |  Tech Stack          (bold navy + italic gray)
    Line 2: date                                  (italic light gray)
    """
    p1 = doc.add_paragraph()
    r1 = p1.add_run(name)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = ACCENT
    r1.font.name = "Arial"
    r2 = p1.add_run(f"  |  {tech_stack}")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    r2.font.name = "Arial"
    set_spacing(p1, before=80, after=0)

    p2 = doc.add_paragraph()
    rd = p2.add_run(date)
    rd.italic = True
    rd.font.size = Pt(9)
    rd.font.color.rgb = LGRAY
    rd.font.name = "Arial"
    set_spacing(p2, before=0, after=10)

    return p1, p2

# ── Load template ──────────────────────────────────────────
doc = Document(TEMPLATE_PATH)

SUMMARY_MARKER    = "{{PROFESSIONAL_SUMMARY}}"
EXPERIENCE_MARKER = "{{WORK_EXPERIENCE}}"
PROJECTS_MARKER   = "{{PROJECTS}}"

def find_marker(doc, marker):
    for i, para in enumerate(doc.paragraphs):
        if marker in para.text:
            return i
    return None

# ── Replace SUMMARY ────────────────────────────────────────
idx = find_marker(doc, SUMMARY_MARKER)
if idx is not None:
    marker_para = doc.paragraphs[idx]
    marker_para.clear()
    run = marker_para.add_run(data["professional_summary"])
    run.font.size      = Pt(9)
    run.font.color.rgb = GRAY
    run.font.name      = "Arial"
    run.italic         = True
    set_spacing(marker_para, before=40, after=60)
else:
    print("WARNING: {{PROFESSIONAL_SUMMARY}} marker not found in template.")

# ── Replace WORK EXPERIENCE ────────────────────────────────
idx = find_marker(doc, EXPERIENCE_MARKER)
if idx is not None:
    anchor_p = doc.paragraphs[idx]._p
    body     = doc.element.body
    new_elements = []

    for job in data["work_experience"]:
        p1, p2 = job_header_para(
            doc,
            job["job_title"],
            job["company"],
            job["city"],
            job["dates"]
        )
        new_elements.append(p1._p)
        new_elements.append(p2._p)
        for b in job["bullets"]:
            new_elements.append(bullet_para(doc, b)._p)

    ref = anchor_p
    for el in new_elements:
        ref.addnext(el)
        ref = el
    body.remove(anchor_p)
else:
    print("WARNING: {{WORK_EXPERIENCE}} marker not found in template.")

# ── Replace PROJECTS ───────────────────────────────────────
idx = find_marker(doc, PROJECTS_MARKER)
if idx is not None:
    anchor_p = doc.paragraphs[idx]._p
    body     = doc.element.body
    new_elements = []

    for proj in data["projects"]:
        p1, p2 = project_header_para(
            doc,
            proj["name"],
            proj["tech_stack"],
            proj["date"]
        )
        new_elements.append(p1._p)
        new_elements.append(p2._p)
        for b in proj["bullets"]:
            new_elements.append(bullet_para(doc, b)._p)

    ref = anchor_p
    for el in new_elements:
        ref.addnext(el)
        ref = el
    body.remove(anchor_p)
else:
    print("WARNING: {{PROJECTS}} marker not found in template.")

# ── Save ───────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Resume saved to: {OUTPUT_PATH}")